import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from reportlab.lib.pagesizes import A4 # type: ignore
from reportlab.pdfgen import canvas # type: ignore
from reportlab.lib.utils import ImageReader # type: ignore
from docx import Document # type: ignore
import json
from datetime import datetime
from tkcalendar import Calendar # type: ignore
from num2words import num2words # type: ignore

def num_para_extenso(valor):
    return num2words(float(valor), lang='pt').replace('-', ' ').capitalize()

def salvar_cliente(nome, cpf, contato, data, horario, valor):
    dados = {"nome": nome, "cpf": cpf, "contato": contato, "data": data, "horario": horario, "valor": valor}
    try:
        with open("clientes.json", "r", encoding="utf-8") as file:
            clientes = json.load(file)
    except FileNotFoundError:
        clientes = []
    clientes.append(dados)
    with open("clientes.json", "w", encoding="utf-8") as file:
        json.dump(clientes, file, indent=4, ensure_ascii=False)

def selecionar_data():
    def confirmar_data():
        entry_data.delete(0, tk.END)
        entry_data.insert(0, cal.get_date())
        popup.destroy()
    
    popup = tk.Toplevel(root)
    popup.title("Selecionar Data")
    cal = Calendar(popup, selectmode='day', year=2025, month=2, day=27)
    cal.pack(pady=20)
    tk.Button(popup, text="Confirmar", command=confirmar_data).pack()

def gerar_recibo():
    nome = entry_nome.get()
    cpf = entry_cpf.get()
    contato = entry_contato.get()
    data = entry_data.get()
    horario = entry_horario.get()
    valor = entry_valor.get()
    
    if not (nome and cpf and contato and data and horario and valor):
        messagebox.showerror("Erro", "Todos os campos devem ser preenchidos!")
        return
    
    valor_extenso = num_para_extenso(valor) + " reais"
    salvar_cliente(nome, cpf, contato, data, horario, valor)
    salvar_pdf(nome, cpf, contato, data, horario, valor, valor_extenso)
    salvar_docx(nome, cpf, contato, data, horario, valor, valor_extenso)
    messagebox.showinfo("Sucesso", "Recibo gerado com sucesso!")

def salvar_pdf(nome, cpf, contato, data, horario, valor, valor_extenso):
    arquivo_pdf = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
    if not arquivo_pdf:
        return
    
    c = canvas.Canvas(arquivo_pdf, pagesize=A4)
    largura, altura = A4
    
    logo_path = "C:/Users/gabri/Downloads/oip.jpeg"
    c.drawImage(ImageReader(logo_path), largura / 2 - 50, altura - 100, width=100, height=100, preserveAspectRatio=True, mask='auto')
    
    for i in range(2):
        y_offset = altura - (i * (altura / 2)) - 150
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(largura / 2, y_offset, "RD PRAG DETETIZAÇÃO")
        c.setFont("Helvetica", 12)
        c.drawCentredString(largura / 2, y_offset - 20, "CNPJ: XX.XXX.XXX/XXXX-XX")
        c.drawString(100, y_offset - 50, f"Nome: {nome}")
        c.drawString(100, y_offset - 70, f"CPF: {cpf}")
        c.drawString(100, y_offset - 90, f"Contato: {contato}")
        c.drawString(100, y_offset - 110, f"Data: {data} - Horário: {horario}")
        c.drawString(100, y_offset - 130, f"Valor: R$ {valor} ({valor_extenso})")
        c.drawString(100, y_offset - 160, "Assinatura: __________________________")
        c.drawString(100, y_offset - 200, "Agradecemos por escolher nossos serviços!")
    
    c.save()

def salvar_docx(nome, cpf, contato, data, horario, valor, valor_extenso):
    arquivo_docx = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if not arquivo_docx:
        return
    
    doc = Document()
    for _ in range(2):
        doc.add_paragraph("RD PRAG DETETIZAÇÃO").bold = True
        doc.add_paragraph("CNPJ: XX.XXX.XXX/XXXX-XX")
        doc.add_paragraph(f"Nome: {nome}")
        doc.add_paragraph(f"CPF: {cpf}")
        doc.add_paragraph(f"Contato: {contato}")
        doc.add_paragraph(f"Data: {data} - Horário: {horario}")
        doc.add_paragraph(f"Valor: R$ {valor} ({valor_extenso})")
        doc.add_paragraph("Assinatura: __________________________")
        doc.add_paragraph("Agradecemos por escolher nossos serviços!")
        doc.add_paragraph("\n" + "-"*50 + "\n")
    
    doc.save(arquivo_docx)

root = tk.Tk()
root.title("Gerador de Recibos")
root.configure(bg="#2E7D32")
root.geometry("400x500")

tk.Label(root, text="Nome:", bg="#2E7D32", fg="white").pack()
entry_nome = tk.Entry(root)
entry_nome.pack()

tk.Label(root, text="CPF:", bg="#2E7D32", fg="white").pack()
entry_cpf = tk.Entry(root)
entry_cpf.pack()

tk.Label(root, text="Contato:", bg="#2E7D32", fg="white").pack()
entry_contato = tk.Entry(root)
entry_contato.pack()

tk.Label(root, text="Data:", bg="#2E7D32", fg="white").pack()
entry_data = tk.Entry(root)
entry_data.pack()
tk.Button(root, text="Selecionar Data", command=selecionar_data).pack()

tk.Label(root, text="Horário:", bg="#2E7D32", fg="white").pack()
entry_horario = tk.Entry(root)
entry_horario.pack()

tk.Label(root, text="Valor (R$):", bg="#2E7D32", fg="white").pack()
entry_valor = tk.Entry(root)
entry_valor.pack()

tk.Button(root, text="Gerar Recibo", command=gerar_recibo, bg="#1B5E20", fg="white").pack(pady=10)


#DESENVOLVIDO POR GABRIEL SANTANA
import webbrowser
def open_link(url):
    webbrowser.open(url)

label = tk.Label(root, text="Desenvolvido por Gabriel Santana | LinkedIn | GitHub", bg="#2E7D32", fg="white")
label.pack(side="bottom", pady=5)

label.bind("<Button-1>", lambda e: open_link("https://www.linkedin.com/in/gabrielsbelarmino/"))
label.bind("<Button-3>", lambda e: open_link("https://github.com/GabrielFSantana"))

root.mainloop()
